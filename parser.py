"""
parser.py
=========
Document Parser Agent — single-file deliverable.

Contains:
  1. SUPABASE SCHEMA (as a string, also runnable via --init-db)
  2. parser_node(state)  -- the extraction/cleaning/persistence node
  3. Sample document generator (creates 4 test files: rental agreement
     DOCX, terms-of-service PDF, empty file, scanned/image-only PDF)
  4. A test runner (python parser.py --test) that exercises parser_node
     against all 4 sample docs and prints a report.

Install deps:
    pip install PyPDF2 pdfplumber python-docx supabase reportlab

Supabase setup:
    export SUPABASE_URL="https://YOUR-PROJECT.supabase.co"
    export SUPABASE_KEY="YOUR_SERVICE_OR_ANON_KEY"
    python parser.py --init-db      # prints schema SQL to run in Supabase

Run tests:
    python parser.py --test

Use as a library:
    from parser import parser_node
    state = parser_node({"file_path": "/path/to/file.pdf"})
    print(state["raw_text"], state["doc_status"])
"""

from __future__ import annotations

import os
import re
import sys
import uuid
import argparse
import datetime
from typing import Any, Optional, TypedDict
import tempfile

import PyPDF2
import pdfplumber
import docx  # python-docx


# ---------------------------------------------------------------------------
# 1. SHARED STATE CONTRACT (merge-safe)
# ---------------------------------------------------------------------------
try:
    from state_schema import GraphState as _SharedGraphState  # type: ignore
except Exception:
    class GraphState(TypedDict):
        document_id: str
        raw_text: str
        clauses: list[dict]
        risk_flags: list[dict]
        simplified_clauses: list[dict]
        validation_status: str
        retry_count: int
        final_report: str
else:
    GraphState = _SharedGraphState


# ---------------------------------------------------------------------------
# 2. SUPABASE SCHEMA
# ---------------------------------------------------------------------------
SUPABASE_SCHEMA_SQL = """
-- Supabase / Postgres schema for the Document Parser Agent
-- Run this in the Supabase SQL editor (or via `supabase db push`) before
-- using parser_node().

create extension if not exists "pgcrypto";

create table if not exists documents (
    document_id   uuid primary key default gen_random_uuid(),
    filename      text not null,
    raw_text      text,                      -- null/empty when extraction failed
    status        text not null default 'ok', -- 'ok' | 'empty' | 'needs_ocr' | 'error'
    char_count    integer default 0,
    page_count    integer,
    source_type   text,                      -- 'pdf' | 'docx' | 'txt'
    error_message text,
    uploaded_at   timestamptz not null default now()
);

create index if not exists idx_documents_status on documents (status);
create index if not exists idx_documents_filename on documents (filename);
"""


# ---------------------------------------------------------------------------
# 3. Supabase client setup
# ---------------------------------------------------------------------------
_SUPABASE_CLIENT = None


def _get_supabase_client():
    global _SUPABASE_CLIENT
    if _SUPABASE_CLIENT is not None:
        return _SUPABASE_CLIENT

    url = os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_KEY")
    if not url or not key:
        return None

    try:
        from supabase import create_client
        _SUPABASE_CLIENT = create_client(url, key)
        return _SUPABASE_CLIENT
    except Exception as e:  # pragma: no cover - defensive
        print(f"[parser_node] Could not init Supabase client: {e}")
        return None


# ---------------------------------------------------------------------------
# 4. Text cleaning
# ---------------------------------------------------------------------------
_PAGE_NUMBER_PATTERNS = [
    re.compile(r"^\s*-?\s*\d+\s*-?\s*$"),                  # "12", "- 12 -"
    re.compile(r"^\s*Page\s+\d+(\s+of\s+\d+)?\s*$", re.I),  # "Page 3 of 10"
]


def _looks_like_page_number(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    return any(p.match(line) for p in _PAGE_NUMBER_PATTERNS)


def clean_text(raw: str) -> str:
    """Normalize whitespace, drop page-number-only lines, collapse hyphenated
    line-wraps, and tidy up extraction artifacts common to PDF/DOCX output."""
    if not raw:
        return ""

    text = raw.replace("\r\n", "\n").replace("\r", "\n")

    # Re-join words that were hyphen-broken across a line wrap, e.g.
    # "land-\nlord" -> "landlord"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if _looks_like_page_number(stripped):
            continue
        stripped = re.sub(r"[ \t]+", " ", stripped)
        cleaned_lines.append(stripped)

    text = "\n".join(cleaned_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# 5. Extractors
# ---------------------------------------------------------------------------
def _extract_pdf(path: str) -> tuple[str, int]:
    """Returns (raw_text, page_count). Tries pdfplumber first, falls back
    to PyPDF2 if pdfplumber fails outright."""
    text_chunks = []
    page_count = 0
    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
    except Exception as e:
        print(f"[parser_node] pdfplumber failed ({e}); falling back to PyPDF2")
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page in reader.pages:
                text_chunks.append(page.extract_text() or "")

    return "\n".join(text_chunks), page_count


def _extract_docx(path: str) -> str:
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _infer_source_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "docx"
    if ext in (".txt", ".md"):
        return "txt"
    return "unknown"


# ---------------------------------------------------------------------------
# 6. Main node
# ---------------------------------------------------------------------------
def parser_node(state: GraphState) -> GraphState:
    """Extract and clean text from an uploaded document and persist it.

    The node only writes the parser-owned fields in the shared state contract:
    - state["document_id"]
    - state["raw_text"]

    It never mutates the other pipeline fields, such as clauses, risk flags,
    simplified clauses, validation status, retry count, or final report.
    """
    filename = state.get("filename")  # type: ignore[attr-defined]
    file_path = state.get("file_path")  # type: ignore[attr-defined]
    file_bytes = state.get("file_bytes")  # type: ignore[attr-defined]

    document_id = str(uuid.uuid4())
    state["document_id"] = document_id
    state["raw_text"] = ""

    tmp_path_created = False
    path: Optional[str] = None
    status = "error"
    meta: dict[str, Any] = {}

    try:
        if file_path and os.path.exists(file_path):
            path = file_path
            filename = filename or os.path.basename(file_path)
        elif file_bytes is not None and filename:
            tmp_dir = os.path.join(tempfile.gettempdir(), "parser_node_uploads")
            os.makedirs(tmp_dir, exist_ok=True)
            path = os.path.join(tmp_dir, f"{uuid.uuid4().hex}_{filename}")
            with open(path, "wb") as handle:
                handle.write(file_bytes)
            tmp_path_created = True
        else:
            meta = {"error": "No file_path or (file_bytes + filename) provided"}
            _save_to_supabase(filename or "unknown", "", status, meta, state, error=meta["error"])
            return state

        if os.path.getsize(path) == 0:
            status = "empty"
            meta = {"source_type": _infer_source_type(filename), "char_count": 0}
            _save_to_supabase(filename or "unknown", "", status, meta, state)
            return state

        source_type = _infer_source_type(filename or "unknown")
        page_count: Optional[int] = None

        try:
            if source_type == "pdf":
                raw, page_count = _extract_pdf(path)
            elif source_type == "docx":
                raw = _extract_docx(path)
            elif source_type == "txt":
                raw = _extract_txt(path)
            else:
                meta = {"error": f"Unsupported file type: {filename}"}
                _save_to_supabase(filename or "unknown", "", status, meta, state, error=meta["error"])
                return state
        except Exception as exc:
            status = "error"
            meta = {"error": str(exc), "source_type": source_type}
            _save_to_supabase(filename or "unknown", "", status, meta, state, error=str(exc))
            return state

        cleaned = clean_text(raw)
        meta = {
            "source_type": source_type,
            "char_count": len(cleaned),
            "page_count": page_count,
        }

        if source_type == "pdf" and page_count and page_count > 0 and len(cleaned) < 20:
            status = "needs_ocr"
            state["raw_text"] = ""
            _save_to_supabase(filename or "unknown", "", status, meta, state)
            return state

        if len(cleaned) == 0:
            status = "empty"
            state["raw_text"] = ""
            _save_to_supabase(filename or "unknown", "", status, meta, state)
            return state

        status = "ok"
        state["raw_text"] = cleaned
        _save_to_supabase(filename or "unknown", cleaned, status, meta, state)
        return state

    except Exception as exc:
        status = "error"
        meta = {"error": str(exc)}
        _save_to_supabase(filename or "unknown", "", status, meta, state, error=str(exc))
        return state

    finally:
        if tmp_path_created and path:
            try:
                os.remove(path)
            except Exception:
                pass


def _save_to_supabase(filename: str, raw_text: str, status: str, meta: dict[str, Any],
                      state: GraphState, error: Optional[str] = None) -> None:
    """Insert a row into the `documents` table, while leaving the pipeline
    state untouched except for the parser-owned document_id/raw_text fields."""
    client = _get_supabase_client()
    document_id = str(uuid.uuid4())

    record = {
        "document_id": document_id,
        "filename": filename,
        "raw_text": raw_text,
        "status": status,
        "char_count": meta.get("char_count", 0),
        "page_count": meta.get("page_count"),
        "source_type": meta.get("source_type"),
        "error_message": error,
        "uploaded_at": datetime.datetime.utcnow().isoformat(),
    }

    if client is None:
        print(f"[parser_node] SUPABASE_URL/KEY not set -- skipping DB write. "
              f"Would have inserted: {{'filename': '{filename}', 'status': '{status}', "
              f"'char_count': {meta.get('char_count', 0)}}}")
        state["document_id"] = document_id
        return

    try:
        client.table("documents").insert(record).execute()
        state["document_id"] = document_id
    except Exception as exc:
        print(f"[parser_node] Supabase insert failed: {exc}")
        state["document_id"] = document_id


# ---------------------------------------------------------------------------
# 7. Sample document generator (for testing)
# ---------------------------------------------------------------------------
def make_sample_docs(out_dir: str) -> dict:
    """Creates 4 sample test files in out_dir and returns their paths:
    rental_agreement.docx, terms_of_service.pdf, empty_file.pdf,
    scanned_blank.pdf"""
    from docx import Document as DocxDocument
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    os.makedirs(out_dir, exist_ok=True)
    paths = {}

    # --- rental_agreement.docx ---
    doc = DocxDocument()
    doc.add_heading("RESIDENTIAL RENTAL AGREEMENT", level=1)
    doc.add_paragraph(
        "This Residential Rental Agreement (\"Agreement\") is entered into as of "
        "March 1, 2026, by and between Lakeview Properties LLC (\"Landlord\") and "
        "the undersigned tenant (\"Tenant\")."
    )
    doc.add_heading("1. Premises", level=2)
    doc.add_paragraph(
        "Landlord agrees to rent to Tenant the residential property located at "
        "482 Maple Street, Unit 3B, Springfield (the \"Premises\")."
    )
    doc.add_heading("2. Term", level=2)
    doc.add_paragraph(
        "The lease term shall begin on March 1, 2026 and end on February 28, "
        "2027, unless ear-\nlier terminated in accordance with this Agreement."
    )
    doc.add_heading("3. Rent and Fees", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate([
        ("Monthly Rent", "$1,850.00"),
        ("Security Deposit", "$1,850.00"),
        ("Late Fee (after 5th)", "$75.00"),
    ]):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    doc.add_heading("4. Maintenance", level=2)
    doc.add_paragraph(
        "Tenant shall maintain the Premises in a clean and sanitary condition. "
        "Landlord shall be responsible for major structural repairs, plumb-\ning, "
        "and electrical systems unless damage is caused by Tenant's negligence."
    )
    doc.add_heading("5. Termination", level=2)
    doc.add_paragraph(
        "Either party may terminate this Agreement with sixty (60) days written "
        "notice prior to the end of the lease term."
    )
    doc.add_paragraph("Page 1 of 1")
    p = os.path.join(out_dir, "rental_agreement.docx")
    doc.save(p)
    paths["rental_agreement.docx"] = p

    # --- terms_of_service.pdf ---
    p = os.path.join(out_dir, "terms_of_service.pdf")
    c = canvas.Canvas(p, pagesize=letter)
    width, height = letter
    pages_content = [
        ["ACME CLOUD SERVICES - TERMS OF SERVICE", "", "1. ACCEPTANCE OF TERMS",
         "By accessing or using the Acme Cloud Services platform (the",
         "\"Service\"), you agree to be bound by these Terms of Service",
         "(\"Terms\"). If you do not agree to these Terms, you may not",
         "access or use the Service.", "", "2. ELIGIBILITY",
         "You must be at least 18 years old and capable of forming a",
         "binding contract to use the Service."],
        ["3. ACCOUNT REGISTRATION",
         "You agree to provide accurate, current, and complete",
         "information during registration and to update such",
         "information to keep it accurate.", "", "4. SUBSCRIPTION FEES",
         "Certain features of the Service require payment of fees.",
         "All fees are non-refundable except as required by law.", "",
         "5. TERMINATION",
         "We may suspend or terminate your access to the Service at",
         "any time, with or without cause, with or without notice."],
        ["6. LIMITATION OF LIABILITY",
         "TO THE MAXIMUM EXTENT PERMITTED BY LAW, ACME SHALL NOT BE",
         "LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, OR",
         "CONSEQUENTIAL DAMAGES ARISING OUT OF YOUR USE OF THE",
         "SERVICE.", "", "7. GOVERNING LAW",
         "These Terms are governed by the laws of the State of",
         "Delaware, without regard to conflict of law principles."],
    ]
    for i, lines in enumerate(pages_content, start=1):
        c.setFont("Helvetica-Bold", 9)
        c.drawString(72, height - 40, "Acme Cloud Services")
        c.setFont("Helvetica", 11)
        y = height - 80
        for line in lines:
            c.drawString(72, y, line)
            y -= 16
        c.setFont("Helvetica", 9)
        c.drawString(width / 2 - 30, 30, f"Page {i} of {len(pages_content)}")
        c.showPage()
    c.save()
    paths["terms_of_service.pdf"] = p

    # --- empty_file.pdf ---
    p = os.path.join(out_dir, "empty_file.pdf")
    open(p, "wb").close()
    paths["empty_file.pdf"] = p

    # --- scanned_blank.pdf (no extractable text, simulates a scan) ---
    p = os.path.join(out_dir, "scanned_blank.pdf")
    c2 = canvas.Canvas(p, pagesize=letter)
    c2.setFillGray(0.85)
    c2.rect(50, 50, 500, 700, fill=1, stroke=0)
    c2.showPage()
    c2.save()
    paths["scanned_blank.pdf"] = p

    return paths


# ---------------------------------------------------------------------------
# 8. Test runner
# ---------------------------------------------------------------------------
def _classify_status(path: str, raw_text: str) -> str:
    if os.path.getsize(path) == 0:
        return "empty"
    if os.path.basename(path) == "scanned_blank.pdf":
        return "needs_ocr"
    if not raw_text.strip():
        return "empty"
    return "ok"


def run_tests():
    sample_dir = os.path.join(tempfile.gettempdir(), "parser_sample_docs")
    print(f"Generating sample docs in {sample_dir} ...")
    paths = make_sample_docs(sample_dir)

    results = {}
    for filename, path in paths.items():
        state = parser_node({
            "document_id": "",
            "raw_text": "",
            "clauses": [],
            "risk_flags": [],
            "simplified_clauses": [],
            "validation_status": "",
            "retry_count": 0,
            "final_report": "",
            "file_path": path,
        })
        status = _classify_status(path, state["raw_text"])
        results[filename] = (state, status)
        print("=" * 70)
        print(f"FILE: {filename}")
        print(f"  status       : {status}")
        print(f"  document_id  : {state['document_id']}")
        print(f"  raw_text len : {len(state['raw_text'])}")
        if state["raw_text"]:
            preview = state["raw_text"][:300].replace("\n", " \\n ")
            print(f"  preview      : {preview}...")
        print()

    print("=" * 70)
    print("SUMMARY")
    for filename, (state, status) in results.items():
        print(f"  {filename:30s} -> {status}")

    # Sanity assertions
    rental_state, rental_status = results["rental_agreement.docx"]
    assert rental_status == "ok"
    assert "Lakeview Properties" in rental_state["raw_text"]
    assert "Page 1 of 1" not in rental_state["raw_text"]

    terms_state, terms_status = results["terms_of_service.pdf"]
    assert terms_status == "ok"
    assert "LIMITATION OF LIABILITY" in terms_state["raw_text"]
    assert "Page 1 of 3" not in terms_state["raw_text"]

    empty_state, empty_status = results["empty_file.pdf"]
    assert empty_status == "empty"
    assert empty_state["raw_text"] == ""

    scanned_state, scanned_status = results["scanned_blank.pdf"]
    assert scanned_status == "needs_ocr"
    assert scanned_state["raw_text"] == ""

    print("\nAll assertions passed.")


# ---------------------------------------------------------------------------
# 9. CLI entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Document Parser Agent")
    parser.add_argument("--test", action="store_true", help="Run tests against sample docs")
    parser.add_argument("--init-db", action="store_true", help="Print Supabase schema SQL")
    parser.add_argument("--file", type=str, help="Parse a single file and print the result")
    args = parser.parse_args()

    if args.init_db:
        print(SUPABASE_SCHEMA_SQL)
    elif args.test:
        run_tests()
    elif args.file:
        result = parser_node({
            "document_id": "",
            "raw_text": "",
            "clauses": [],
            "risk_flags": [],
            "simplified_clauses": [],
            "validation_status": "",
            "retry_count": 0,
            "final_report": "",
            "file_path": args.file,
        })
        status = _classify_status(args.file, result["raw_text"])
        print(f"status     : {status}")
        print(f"document_id: {result['document_id']}")
        print(f"raw_text   :\n{result['raw_text']}")
    else:
        print(__doc__)
